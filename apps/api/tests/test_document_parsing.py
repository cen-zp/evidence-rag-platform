from docx import Document as DocxDocument

from app.services.document_parsing import parse_document


def test_parse_docx_extracts_non_empty_paragraphs(tmp_path) -> None:
    path = tmp_path / "handbook.docx"
    document = DocxDocument()
    document.add_heading("Release process", level=1)
    document.add_paragraph("Verify the change before publishing.")
    document.save(path)

    pages = parse_document(
        path,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    assert len(pages) == 1
    assert pages[0].page_number is None
    assert pages[0].text == "Release process\nVerify the change before publishing."
